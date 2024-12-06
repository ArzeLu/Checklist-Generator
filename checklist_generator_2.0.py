# TODO: replace sys exit warnings with pop up windows.

from docx import Document # type: ignore (warning suppressant for vscode that doesn't have the library globally)
from docx.shared import Pt # type: ignore (warning suppressant for vscode that doesn't have the library globally)
from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
import pytz # type: ignore
import sys
import os
import time
import random
import re

APP_VERSION = "2.0.0"

class Helper():
    def __init__(self):
        self.taiwan_timezone = pytz.timezone('Asia/Taipei')

    ## Turns a raw time string of, for example: "2024-12-11 23:08:17",
    ## from taiwan timezone to the US time,
    ## then format it in the checklist format: "%Y-%m-%d / %I:%M %p" ("2024-12-11 11:08 PM")
    def convert_timezone(self, time):
        us_time = datetime.fromisoformat(time)
        us_time = us_time.astimezone(self.taiwan_timezone)
        us_time = us_time.strftime(r"%Y-%m-%d / %I:%M %p")

        return us_time

class TextFileHandler():
    def __init__(self):
        self.helper = Helper()
        self.source_files_dir = ""

    def get_file_names(self, source_files_dir):
        self.source_files_dir = source_files_dir
        source_file_names = []
        clean_file_names = []

        # Get burn in file names 
        for root, dirs, files in os.walk(source_files_dir):
            source_file_names.extend(files)

        # Remove the file extensions from the file name strings
        for file in source_file_names:
            clean_file_names.append(os.path.splitext(file)[0])

        return clean_file_names

    def find_info(self, file_name, source_file_dir):
        serial = file_name
        start_time = ""
        end_time = ""

        # find burn in times from the burn in .txt
        dates = []
        with open(source_file_dir + "/" + file_name + ".log", encoding = "utf_16") as file:
            lines = file.readlines()

            for line in lines:
                line = line.rstrip()
                date_match = re.search(r"\d{4}[-]\d{2}[-]\d{2}\s\d{2}[:]\d{2}[:]\d{2}", line)
                if date_match is not None:
                    dates.append(date_match.group())

        start_time = None
        end_time = None

        try:
            # there will be three matches, only get the first and last one.
            # look at a burn-in report for reference
            start_time = dates[0]
            end_time = dates[2]

            start_time = self.helper.convert_timezone(start_time)
            end_time = self.helper.convert_timezone(end_time)

        except:
            print("======")
            print(dates)
            print("======")
            sys.exit(f"""
            =============== Error! =================\n
            File: "{file_name}" doesn't have enough dates\n
            ========================================\n
            """)

        return [serial, start_time, end_time]
        

class DocxHandler():
    def __init__(self):
        self.txt_handler = TextFileHandler()
        self.serial_placeholder = "{serial number}"
        self.start_time_placeholder = "{start time}"
        self.end_time_placeholder = "{end time}"

    ## From source file to the checklist:
    ## Fill in serial number
    ## Fill in burn in times
    ## Checks if any placeholder is missing from the checklist and throw appropriate exceptions
    ##
    ## Arguments: 1. The checklist template, 2. One target file name
    def fill_checklist_info(self, checklist, file_name, source_file_dir):
        new_checklist = Document()
        paragraphs = checklist.paragraphs
        tables = checklist.tables
        # fill in info
        serial, start_time, end_time = self.txt_handler.find_info(file_name, source_file_dir)
        serial_placeholder_found = False
        start_time_placeholder_found = False
        end_time_placeholder_found = False

        for table in tables:
            rows = table.rows
            new_table = new_checklist.add_table()
            for row in rows:
                cells = row.cells
                for cell in cells:
                    text = cell.text
                    print(text)
                    if self.serial_placeholder in text:
                        text.replace(self.serial_placeholder, serial)
                        serial_placeholder_found = True
                new_table.add_row(row)

        for paragraph in paragraphs:
            text = paragraph.text
            if self.start_time_placeholder in text:
                text.replace(self.start_time_placeholder, start_time)
                start_time_placeholder_found = True
                new_checklist.add_paragraph(text)
                continue

            if self.end_time_placeholder in text:
                text.replace(self.end_time_placeholder, end_time)
                end_time_placeholder_found = True
                new_checklist.add_paragraph(text)
                continue

            new_checklist.add_paragraph(text)

        # Check for any missing placeholders
        error = False

        if not serial_placeholder_found:
            print("""
            =============== Error! =================\n
            Checklist template doesn't have {serial number} placeholder!\n
            ========================================\n
            """)
            error = True

        if not start_time_placeholder_found:
            print("""
            =============== Error! =================\n
            Checklist template doesn't have {start time} placeholder!\n
            ========================================\n
            """)
            error = True

        if not end_time_placeholder_found:
            print("""
            =============== Error! =================\n
            Checklist template doesn't have {end time} placeholder!\n
            ========================================\n
            """)
            error = True

        if error:
            sys.exit()

        return new_checklist

    ## Generate checklists with appropriate fields
    ## Arguments: 1. directory of the checklist, 2. directory of the source files, 3. directory of generated files
    ## Only checking the validity of directories here because it's the only time when they're used.
    def generate_checklists(self, checklist_dir, source_files_dir, destination_dir):
        source_file_names = self.txt_handler.get_file_names(source_files_dir)
        checklist = Document(checklist_dir)

        # check if the directories are valid.
        # use exists for both files and directories
        error = False
        if os.path.exists(checklist_dir) is not True:
            print("""
            =============== Error! =================\n
            Invalid checklist directory!\n
            ========================================\n
            """)
            error = True

        if os.path.isdir(source_files_dir) is not True:
            print("""
            =============== Error! =================
            Invalid source files directory!\n
            ========================================\n
            """)
            error = True

        if os.path.isdir(destination_dir) is not True:
            print("""
            =============== Error! =================\n
            Invalid destination directory!\n
            ========================================\n
            """)
            error = True

        if error:
            sys.exit()
        
        for file_name in source_file_names:
            new_checklist = self.fill_checklist_info(checklist, file_name, source_files_dir)
            new_checklist.save(destination_dir + "/" + file_name + ".docx")

        

class UserInterface():
    def __init__(self):
        self.root = Tk()
        self.doc_handler = DocxHandler()
        self.source_files_dir = StringVar(self.root)
        self.checklist_template_dir = StringVar(self.root)
        self.destination_dir = StringVar(self.root)
        self.directory_entries = [self.source_files_dir, self.checklist_template_dir, self.destination_dir]
        self.emotes = ['📁', '👽', '🐱', '🦷', '🎁', '👾', '🦜', '💀', '💩', '🍥', '🥶', '💅', '🍕', '🚗', '🎃']
        self.emotes_count = 15
    
    def select_burnin_dir(self):
        self.source_files_dir.set(filedialog.askdirectory(initialdir = "/", title = "🔥🔥🔥🔥🔥"))

    def select_checklist_dir(self):
        self.checklist_template_dir.set(filedialog.askopenfilename(initialdir = "/", title = "📄📄📄📄📄", filetypes = [("💋💋💋💋💋💋💋💋", "*.*")]))

    def select_destination_dir(self):
        self.destination_dir.set(filedialog.askdirectory(initialdir = "/", title = "🎯🎯🎯🎯🎯"))

    def choose_emote(self):
        n = random.randrange(0, self.emotes_count)  # first argument is included, not the second.
        return self.emotes[n]

    def generate_button_action(self):
        self.doc_handler.generate_checklists(self.checklist_template_dir.get(), self.source_files_dir.get(), self.destination_dir.get())

    def reroll(self):
        for entry in self.directory_entries:
            entry.set("")
	
        self.root.update()  # This bypasses the buffer for the gui
        time.sleep(0.3)

        for entry in self.directory_entries:
            entry.set(self.choose_emote())
            self.root.update()
            time.sleep(0.2)

    def run(self):
        self.root.title("Checklist Generator")
        self.root.geometry("600x200")
        self.root.columnconfigure(0, weight = 1)
        self.root.rowconfigure(0, weight = 1)
        self.root.rowconfigure(1, weight = 1)
        self.root.resizable(False, False)

        s = ttk.Style()

        ## Elements of the first three rows ##
        s.configure("Top.TFrame")
        top_frame = ttk.Frame(self.root, padding = (10, 20, 5, 5), style = "Top.TFrame")
        top_frame.grid(column = 0, row = 0, sticky = N+S+W+E)  # set the starting position of the top_frame

        burnin_label = ttk.Label(top_frame, text = "Burn-in files: ", font = (15))
        checklist_label = ttk.Label(top_frame, text = "Checklist template: ", font = (15))
        destination_label = ttk.Label(top_frame, text = "Destination folder:", font = (15))

        burnin_dir = ttk.Entry(top_frame, textvariable = self.source_files_dir, width = 25, background = "white", font = (15))
        checklist_dir = ttk.Entry(top_frame, textvariable = self.checklist_template_dir, width = 25, background = "white", font = (15))
        destination_dir = ttk.Entry(top_frame, textvariable = self.destination_dir, width = 25, background = "white", font = (15))

        burnin_button = ttk.Button(top_frame, text = "Browse", command = self.select_burnin_dir)
        checklist_button = ttk.Button(top_frame, text = "Browse", command = self.select_checklist_dir)
        destination_button = ttk.Button(top_frame, text = "Browse", command = self.select_destination_dir)

        burnin_label.grid(column = 0, row = 0, columnspan = 3, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)
        checklist_label.grid(column = 0, row = 1, columnspan = 3, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)
        destination_label.grid(column = 0, row = 2, columnspan = 3, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)

        random.seed(time.time_ns())  # time_ns() instead of time() because nanoseconds better
        self.reroll()
        burnin_dir.grid(column = 3, row = 0, padx = (0, 1), sticky = N+S+W+E)
        checklist_dir.grid(column = 3, row = 1, padx = (1, 1), sticky = N+S+W+E)
        destination_dir.grid(column = 3, row = 2, padx = (1, 1), sticky = N+S+W+E)

        burnin_button.grid(column = 4, row = 0, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)
        checklist_button.grid(column = 4, row = 1, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)
        destination_button.grid(column = 4, row = 2, padx = (5, 5), pady = (0, 5), sticky = N+S+W+E)

        ## Remaining buttons on the bottom half of the screen ##
        s.configure("Bottom.TFrame")
        bottom_frame = ttk.Frame(self.root, padding = (5, 5, 5, 5), style = "Bottom.TFrame")
        bottom_frame.grid(column = 0, row = 1, sticky = N+S+W+E)

        # Function buttons
        s.configure("Button.TButton", font = (15))

        insert_button = ttk.Button(bottom_frame, text = "Insert Serial Numbers", style = "Button.TButton")
        insert_button.grid(column = 0, row = 0, padx = (10, 10), pady = (10, 10), sticky = N+S+W+E)

        generate_button = ttk.Button(bottom_frame, text = "Generate Checklists", style = "Button.TButton", command = self.generate_button_action)
        generate_button.grid(column = 1, row = 0, padx = (10, 10), pady = (10, 10), sticky = N+S+W+E)    

        gamble_button = ttk.Button(bottom_frame, text = "Reroll", style = "Button.TButton", command = self.reroll)
        gamble_button.grid(column = 2, row = 0, padx = (10, 10), pady = (10, 10), sticky = N+S+W+E)

        self.root.mainloop()

class Main():
    def __init__(self):
        self.gui = UserInterface()
    def run(self):
        self.gui.run()
        
if __name__ == "__main__":
    driver = Main()
    driver.run()